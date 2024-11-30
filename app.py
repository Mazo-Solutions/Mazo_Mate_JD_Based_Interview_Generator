import streamlit as st
import pandas as pd
import os
from dotenv import load_dotenv
import google.generativeai as genai
import io
from docx import Document
import PyPDF2
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from io import BytesIO

# Load environment variables from .env file
load_dotenv()

# Configure Gemini API key
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# Function to parse uploaded JD or Resume file (PDF or Word)
def parse_uploaded_file(uploaded_file):
    """Parse the text from an uploaded file (JD or Resume)."""
    try:
        if uploaded_file.name.endswith(".pdf"):
            # Extract text from PDF
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text = "".join(page.extract_text() for page in pdf_reader.pages)
        elif uploaded_file.name.endswith(".docx"):
            # Extract text from Word document
            doc = Document(uploaded_file)
            text = "\n".join(para.text for para in doc.paragraphs)
        else:
            st.error("Unsupported file format. Please upload a PDF or Word document.")
            return None
        return text.strip()
    except Exception as e:
        st.error(f"Error parsing the uploaded file: {e}")
        return None

# Function to generate interview questions using Gemini
def generate_interview_questions_from_jd(jd_text, num_questions):
    """Generate interview questions and answers using JD content."""
    try:
        # Define generation configuration
        generation_config = {
            "temperature": 1,
            "top_p": 0.95,
            "top_k": 40,
            "max_output_tokens": 8192,
            "response_mime_type": "text/plain",
        }

        # Initialize the Gemini model
        model = genai.GenerativeModel(
            model_name="gemini-1.5-flash",
            generation_config=generation_config,
        )

        # Start a chat session
        chat_session = model.start_chat()

        # Build the prompt dynamically based on the JD text
        prompt = (
            f"Based on the following job description, generate {num_questions} interview questions and answers. "
            f"The questions should be relevant to the job requirements:\n\n"
            f"{jd_text}"
        )

        # Send the prompt and get the response
        response = chat_session.send_message(prompt)
        return response.text.strip()
    except Exception as e:
        st.error(f"Error generating questions: {e}")
        return ""

# Function to compute match score between JD and Resume
def compute_match_score(jd_text, resume_text):
    """Compute the similarity score between the JD and resume text."""
    try:
        # Vectorize the texts
        vectorizer = TfidfVectorizer()
        tfidf_matrix = vectorizer.fit_transform([jd_text, resume_text])
        
        # Calculate cosine similarity
        similarity_score = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
        
        # Convert to percentage
        match_score = round(similarity_score * 100, 2)
        return match_score
    except Exception as e:
        st.error(f"Error calculating match score: {e}")
        return None

# Function to export data to an Excel file and return the file as a BytesIO object
def export_to_excel(data):
    """Exports the data to an Excel file and returns the file in memory as a BytesIO object."""
    try:
        # Create a pandas DataFrame from the list of dictionaries
        df = pd.DataFrame(data)

        # Save to a BytesIO buffer using openpyxl (default engine for pandas)
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Interview Questions')

        output.seek(0)  # Rewind the buffer
        return output
    except Exception as e:
        st.error(f"Error exporting to Excel: {e}")
        return None

# Function to export data to a Word document and return the file as a BytesIO object
def export_to_word(data):
    """Exports the data to a Word document and returns the file as a BytesIO object."""
    try:
        # Create a Word document
        doc = Document()
        doc.add_heading('Interview Questions and Answers', 0)

        for qa in data:
            # Add question and answer to Word
            doc.add_heading(qa['Question'], level=1)
            doc.add_paragraph(qa['Answer'])

        # Save to a BytesIO buffer
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)  # Rewind the buffer
        return output
    except Exception as e:
        st.error(f"Error exporting to Word: {e}")
        return None

# Main Streamlit app
def main():
    # Add logo at the center
    logo_url = "https://mazobeam.com/wp-content/uploads/2023/12/mazoid-1.png"
    st.markdown(
        f"""
        <div style="text-align: center;">
            <img src="{logo_url}" alt="Logo" style="width: 300px; margin-bottom: 20px;">
        </div>
        """,
        unsafe_allow_html=True
    )

    st.title("MazoMate - JD-Based Interview Question Generator")

    # File upload for JD
    uploaded_jd_file = st.file_uploader("Upload Job Description (PDF or Word)", type=["pdf", "docx"], key="jd_file")

    if uploaded_jd_file:
        st.info("Parsing the uploaded Job Description...")
        jd_text = parse_uploaded_file(uploaded_jd_file)

        if jd_text:
            st.success("Job Description parsed successfully!")
            st.text_area("Parsed JD Content", jd_text, height=300)

            # Number of questions to generate
            num_questions = st.number_input("Number of Questions to Generate", min_value=1, max_value=100, step=1, value=10)

            # Generate Questions Button
            if st.button("Generate Questions"):
                st.info("Generating interview questions. Please wait...")
                generated_content = generate_interview_questions_from_jd(jd_text, num_questions)

                if generated_content:
                    st.success("Questions generated successfully!")
                    st.write("Generated Questions and Answers")
                    st.text_area("Questions & Answers", generated_content, height=300)

                    # Parse questions and answers
                    qa_pairs = []
                    lines = generated_content.split('\n')
                    for i in range(0, len(lines), 2):  # Assuming questions and answers alternate
                        question = lines[i].strip() if i < len(lines) else ""
                        answer = lines[i + 1].strip() if (i + 1) < len(lines) else ""
                        qa_pairs.append({"Question": question, "Answer": answer})

                    # Export to Excel
                    excel_file = export_to_excel(qa_pairs)
                    if excel_file:
                        st.download_button(
                            label="Download as Excel",
                            data=excel_file,
                            file_name="Mazo_Interview_Questions.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                        )

                    # Export to Word
                    word_file = export_to_word(qa_pairs)
                    if word_file:
                        st.download_button(
                            label="Download as Word",
                            data=word_file,
                            file_name="Mazo_Interview_Questions.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

if __name__ == "__main__":
    main()
